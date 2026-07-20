from app.ingest.extractor import extract


def test_text_extractor_falls_back_to_gb18030(tmp_path):
    source = tmp_path / "题库.txt"
    source.write_bytes("第一题：什么是事务？".encode("gb18030"))

    result = extract(str(source))

    assert result.full_text == "第一题：什么是事务？"


def test_docx_extractor_preserves_heading_levels(tmp_path):
    from docx import Document

    source = tmp_path / "题库.docx"
    doc = Document()
    doc.add_heading("Redis 面试题", level=2)
    doc.add_heading("题1｜Redis 为什么快？", level=3)
    doc.add_paragraph("答案：主要操作在内存中完成。")
    doc.save(source)

    result = extract(str(source))

    assert "## Redis 面试题" in result.full_text
    assert "### 题1｜Redis 为什么快？" in result.full_text
